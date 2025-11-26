#!/usr/bin/env python3
"""
Prospectus Compliance Analyzer
Replicates the exact analytical approach for document compliance verification
"""

import json
import sys
import os
from pathlib import Path
from typing import Dict, List, Any, Tuple
import httpx
from openai import OpenAI
import docx
import tiktoken
from dotenv import load_dotenv
import re

# Load environment variables from .env file
load_dotenv()


class ProspectusComplianceAnalyzer:
    """
    Analyzes fund presentation documents against prospectus rules
    using the same 6-phase approach
    """
    
    def __init__(self, api_key: str):
        """Initialize the analyzer with API credentials"""
        self.api_key = api_key
        self.base_url = "https://tokenfactory.esprit.tn/api"
        self.model = "hosted_vllm/Llama-3.1-70B-Instruct"
        self.max_tokens_per_chunk = 19500
        
        # Initialize OpenAI client with SSL verification disabled
        http_client = httpx.Client(verify=False)
        self.client = OpenAI(
            api_key=api_key,
            base_url=self.base_url,
            http_client=http_client
        )
        
        # Initialize tokenizer for chunking
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        
        # Storage for analysis results
        self.document_data = {}
        self.rules_data = {}
        self.metadata = {}
        self.prospectus_text = ""
        self.violations = []
        self.compliant_rules = []
        
    def count_tokens(self, text: str) -> int:
        """Count tokens in text"""
        return len(self.tokenizer.encode(text))
    
    def chunk_text(self, text: str, max_tokens: int = 19500) -> List[str]:
        """Split text into chunks that fit within token limit"""
        tokens = self.tokenizer.encode(text)
        chunks = []
        
        for i in range(0, len(tokens), max_tokens):
            chunk_tokens = tokens[i:i + max_tokens]
            chunk_text = self.tokenizer.decode(chunk_tokens)
            chunks.append(chunk_text)
        
        return chunks
    
    def load_json_file(self, filepath: str) -> Dict:
        """Load and parse JSON file"""
        print(f"üìÑ Loading {filepath}...")
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def load_docx_file(self, filepath: str) -> str:
        """Extract text from DOCX file"""
        print(f"üìÑ Loading {filepath}...")
        doc = docx.Document(filepath)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        text = '\n'.join(full_text)
        print(f"   ‚úì Extracted {len(text)} characters, {self.count_tokens(text)} tokens")
        return text
    
    def call_llm(self, system_prompt: str, user_prompt: str, temperature: float = 0.3) -> str:
        """Call the LLM API with given prompts"""
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=temperature,
                max_tokens=2000,
                top_p=0.9,
                frequency_penalty=0.0,
                presence_penalty=0.0
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"‚ùå Error calling LLM: {e}")
            return ""
    
    # ========================================================================
    # PHASE 1: DOCUMENT UNDERSTANDING (STRUCTURAL ANALYSIS)
    # ========================================================================
    
    def phase1_document_understanding(self):
        """
        Phase 1: Analyze document structure and map content organization
        """
        print("\n" + "="*80)
        print("PHASE 1: DOCUMENT UNDERSTANDING (STRUCTURAL ANALYSIS)")
        print("="*80)
        
        system_prompt = """You are a compliance analyst performing structural analysis of a fund presentation document.
Your task is to map the document structure and identify key data fields."""
        
        user_prompt = f"""Analyze this fund presentation document structure:

{json.dumps(self.document_data, indent=2)[:5000]}

Provide a concise structural analysis covering:
1. Document type and organization (slides/sections)
2. Key fields present (fund name, benchmark, fees, risks, etc.)
3. Key fields MISSING or EMPTY (look for "", [], null values)
4. Where critical information is located (which slides)
5. Any immediate red flags (empty required fields)

Format your response as a structured analysis."""
        
        analysis = self.call_llm(system_prompt, user_prompt)
        print("\nüìä Structural Analysis:")
        print(analysis)
        
        return {
            "phase": "Phase 1: Document Understanding",
            "analysis": analysis
        }
    
    # ========================================================================
    # PHASE 2: RULES FRAMEWORK ANALYSIS
    # ========================================================================
    
    def phase2_rules_framework(self):
        """
        Phase 2: Analyze rules framework to understand requirements
        """
        print("\n" + "="*80)
        print("PHASE 2: RULES FRAMEWORK ANALYSIS")
        print("="*80)
        
        system_prompt = """You are a compliance analyst analyzing regulatory rules.
Your task is to understand the rules framework and categorize requirements."""
        
        user_prompt = f"""Analyze this prospectus rules framework:

{json.dumps(self.rules_data, indent=2)}

Provide analysis covering:
1. Total number of rules and their severity distribution (critical/major)
2. Validation types used (exact_match, semantic_match, range_check, etc.)
3. Most critical rules for non-professional clients
4. Fields that multiple rules check (high-priority fields)
5. Rules that require prospectus document for verification

Format as a structured summary."""
        
        analysis = self.call_llm(system_prompt, user_prompt)
        print("\nüìã Rules Framework Analysis:")
        print(analysis)
        
        return {
            "phase": "Phase 2: Rules Framework",
            "analysis": analysis
        }
    
    # ========================================================================
    # PHASE 3: FIELD PRESENCE CHECK (QUICK SCAN)
    # ========================================================================
    
    def phase3_field_presence_check(self):
        """
        Phase 3: Check if required fields exist in the document
        """
        print("\n" + "="*80)
        print("PHASE 3: FIELD PRESENCE CHECK (QUICK SCAN)")
        print("="*80)
        
        findings = []
        
        for rule in self.rules_data.get("rules", []):
            rule_id = rule["rule_id"]
            fields_to_check = rule["fields_to_check"]
            severity = rule["severity"]
            
            print(f"\nüîç Checking {rule_id}: {rule['rule_text']}")
            print(f"   Fields to check: {fields_to_check}")
            
            system_prompt = """You are a compliance analyst checking field presence.
Your task is to determine if required fields exist and contain data."""
            
            user_prompt = f"""Check if these fields exist and contain data in the document:

Rule: {rule_id} - {rule['rule_text']}
Severity: {severity}
Fields to check: {fields_to_check}

Document excerpt:
{json.dumps(self.document_data, indent=2)[:3000]}

For each field, respond with:
- FOUND: Field exists with data
- EMPTY: Field exists but is empty ("", [], null)
- MISSING: Field does not exist

Format: Field_name: STATUS - Brief explanation"""
            
            result = self.call_llm(system_prompt, user_prompt, temperature=0.1)
            print(f"   Result: {result[:200]}...")
            
            findings.append({
                "rule_id": rule_id,
                "severity": severity,
                "fields_checked": fields_to_check,
                "result": result
            })
        
        return {
            "phase": "Phase 3: Field Presence Check",
            "findings": findings
        }
    
    # ========================================================================
    # PHASE 4: CONTENT QUALITY ANALYSIS (DEEP DIVE)
    # ========================================================================
    
    def phase4_content_quality_analysis(self):
        """
        Phase 4: Evaluate quality and completeness of existing content
        """
        print("\n" + "="*80)
        print("PHASE 4: CONTENT QUALITY ANALYSIS (DEEP DIVE)")
        print("="*80)
        
        system_prompt = """You are a compliance analyst evaluating content quality.
Your task is to assess if existing content is complete, clear, and unambiguous."""
        
        # Focus on key content areas
        areas_to_check = [
            {
                "area": "Risk Disclosure",
                "location": "slide_2 and page_de_garde",
                "criteria": "Complete risk enumeration, SRI indicator, clear disclaimers"
            },
            {
                "area": "Investment Strategy",
                "location": "page_de_garde and pages_suivantes",
                "criteria": "Asset allocation thresholds, geographic allocation, strategy description"
            },
            {
                "area": "Benchmark Specification",
                "location": "page_de_fin.fund_characteristics",
                "criteria": "Exact benchmark name, dividend treatment, index specification"
            },
            {
                "area": "Fund Characteristics",
                "location": "page_de_fin",
                "criteria": "Fees, minimum investment, AUM, all required fields"
            }
        ]
        
        findings = []
        
        for area in areas_to_check:
            print(f"\nüî¨ Analyzing: {area['area']}")
            
            user_prompt = f"""Evaluate the quality and completeness of this content area:

Area: {area['area']}
Location: {area['location']}
Quality Criteria: {area['criteria']}

Document data:
{json.dumps(self.document_data, indent=2)[:4000]}

Assess:
1. Is the content present?
2. Is it complete and detailed?
3. Is it clear and unambiguous?
4. Are there any quality issues (vague wording, missing details)?
5. What improvements are needed?

Provide a quality rating: EXCELLENT / GOOD / PARTIAL / POOR / MISSING"""
            
            result = self.call_llm(system_prompt, user_prompt)
            print(f"   {result[:150]}...")
            
            findings.append({
                "area": area['area'],
                "quality_assessment": result
            })
        
        return {
            "phase": "Phase 4: Content Quality Analysis",
            "findings": findings
        }
    
    # ========================================================================
    # PHASE 5: CROSS-REFERENCE ANALYSIS
    # ========================================================================
    
    def phase5_cross_reference_analysis(self):
        """
        Phase 5: Check for internal consistency across document
        """
        print("\n" + "="*80)
        print("PHASE 5: CROSS-REFERENCE ANALYSIS")
        print("="*80)
        
        system_prompt = """You are a compliance analyst checking internal consistency.
Your task is to find contradictions or inconsistencies within the document."""
        
        user_prompt = f"""Check for internal consistency in this document:

{json.dumps(self.document_data, indent=2)[:5000]}

Look for:
1. Risks mentioned in one place but not listed in another
2. Conflicting information across slides
3. Data mentioned but not in final characteristics table
4. Inconsistent terminology or naming
5. Missing cross-references

List any inconsistencies found with specific locations."""
        
        result = self.call_llm(system_prompt, user_prompt)
        print("\nüîÑ Cross-Reference Analysis:")
        print(result)
        
        return {
            "phase": "Phase 5: Cross-Reference Analysis",
            "findings": result
        }
    
    # ========================================================================
    # PHASE 6: CONTEXT-AWARE EVALUATION
    # ========================================================================
    
    def phase6_context_aware_evaluation(self):
        """
        Phase 6: Apply context from metadata (client type, product type)
        """
        print("\n" + "="*80)
        print("PHASE 6: CONTEXT-AWARE EVALUATION")
        print("="*80)
        
        system_prompt = """You are a compliance analyst applying regulatory context.
Your task is to adjust requirements based on client type and product characteristics."""
        
        user_prompt = f"""Apply regulatory context to this analysis:

Metadata:
{json.dumps(self.metadata, indent=2)}

Document Type: {self.document_data.get('document_metadata', {}).get('document_type')}
Client Type: {'Non-Professional' if not self.metadata.get('Le client est-il un professionnel') else 'Professional'}
New Product: {self.metadata.get('Le document fait-il r√©f√©rence √† un nouveau Produit')}

Based on this context:
1. Which rules become MORE critical? (e.g., SRI for retail clients)
2. Which disclosures are MANDATORY vs OPTIONAL?
3. What additional protections are required?
4. How does this affect severity of violations?

Provide context-adjusted priorities."""
        
        result = self.call_llm(system_prompt, user_prompt)
        print("\nüéØ Context-Aware Evaluation:")
        print(result)
        
        return {
            "phase": "Phase 6: Context-Aware Evaluation",
            "context_analysis": result
        }
    
    # ========================================================================
    # PHASE 7: PROSPECTUS VERIFICATION (WITH CHUNKING)
    # ========================================================================
    
    def phase7_prospectus_verification(self):
        """
        Phase 7: Verify document content against prospectus (chunked)
        """
        print("\n" + "="*80)
        print("PHASE 7: PROSPECTUS VERIFICATION (CHUNKED ANALYSIS)")
        print("="*80)
        
        if not self.prospectus_text:
            print("‚ö†Ô∏è  No prospectus document loaded, skipping verification")
            return {"phase": "Phase 7: Prospectus Verification", "status": "skipped"}
        
        # Chunk the prospectus
        chunks = self.chunk_text(self.prospectus_text, self.max_tokens_per_chunk)
        print(f"üìë Prospectus split into {len(chunks)} chunks")
        
        # Extract key information from each chunk
        prospectus_data = {
            "sri_rating": None,
            "risk_list": [],
            "asset_allocation": None,
            "benchmark": None,
            "minimum_investment": None,
            "fees": None,
            "investment_objective": None
        }
        
        system_prompt = """You are a compliance analyst extracting key information from a prospectus.
Extract specific regulatory data points accurately."""
        
        for i, chunk in enumerate(chunks):
            print(f"\nüîç Analyzing chunk {i+1}/{len(chunks)}...")
            
            user_prompt = f"""Extract these key data points from this prospectus section:

1. SRI/SRRI rating (e.g., "5/7", "6/7")
2. Complete list of risks mentioned
3. Asset allocation thresholds (e.g., "80-100% equities")
4. Benchmark name and specification
5. Minimum investment amount
6. Management fees / TER
7. Investment objective statement

Prospectus section:
{chunk[:15000]}

Format as JSON with keys: sri_rating, risks, asset_allocation, benchmark, minimum_investment, fees, objective
If not found in this section, use null."""
            
            result = self.call_llm(system_prompt, user_prompt, temperature=0.1)
            
            # Try to parse JSON response with robust extraction
            try:
                # Try to extract JSON from the response (might have extra text)
                json_start = result.find('{')
                json_end = result.rfind('}') + 1
                
                if json_start != -1 and json_end > json_start:
                    json_str = result[json_start:json_end]
                    chunk_data = json.loads(json_str)
                    
                    # Merge data
                    for key, value in chunk_data.items():
                        if value and value != "null" and value != None:
                            if key == "risks" and isinstance(value, list):
                                prospectus_data["risk_list"].extend(value)
                            elif not prospectus_data.get(key):
                                prospectus_data[key] = value
                    
                    print(f"   ‚úì Extracted data from chunk {i+1}")
                else:
                    print(f"   ‚ö†Ô∏è  No JSON found in chunk {i+1}")
            except json.JSONDecodeError as e:
                print(f"   ‚ö†Ô∏è  Could not parse JSON from chunk {i+1}: {e}")
                # Continue anyway - we might get data from other chunks
            except Exception as e:
                print(f"   ‚ö†Ô∏è  Error processing chunk {i+1}: {e}")
        
        print("\nüìä Extracted Prospectus Data:")
        print(json.dumps(prospectus_data, indent=2))
        
        # Now compare with document
        print("\nüî¨ Comparing document vs prospectus...")
        
        system_prompt = """You are a compliance analyst comparing document against prospectus.
Identify exact matches, mismatches, and missing information."""
        
        user_prompt = f"""Compare the presentation document against prospectus data:

PROSPECTUS DATA:
{json.dumps(prospectus_data, indent=2)}

PRESENTATION DOCUMENT:
{json.dumps(self.document_data, indent=2)[:5000]}

For each key field, determine:
- MATCH: Content matches prospectus
- MISMATCH: Content differs from prospectus
- MISSING: Required field not in presentation
- INCOMPLETE: Partial information provided

List all discrepancies with severity (CRITICAL/MAJOR/MINOR)."""
        
        comparison = self.call_llm(system_prompt, user_prompt)
        print(comparison)
        
        return {
            "phase": "Phase 7: Prospectus Verification",
            "prospectus_data": prospectus_data,
            "comparison": comparison
        }
    
    # ========================================================================
    # FINAL SYNTHESIS: GENERATE COMPLIANCE REPORT
    # ========================================================================
    
    def generate_final_report(self, phase_results: List[Dict]) -> Dict:
        """
        Synthesize all phase results into final compliance report
        """
        print("\n" + "="*80)
        print("GENERATING FINAL COMPLIANCE REPORT")
        print("="*80)
        
        system_prompt = """You are a senior compliance analyst generating a final compliance report.
Synthesize all analysis phases into a comprehensive, actionable report."""
        
        user_prompt = f"""Generate a comprehensive compliance report based on these analysis phases:

PHASE RESULTS:
{json.dumps(phase_results, indent=2)[:15000]}

RULES FRAMEWORK:
{json.dumps(self.rules_data, indent=2)[:3000]}

METADATA:
{json.dumps(self.metadata, indent=2)}

Generate a report with:

1. EXECUTIVE SUMMARY
   - Overall compliance status
   - Critical/Major/Minor issue counts
   - Distribution recommendation

2. CRITICAL VIOLATIONS (list each with):
   - Rule ID and description
   - Specific finding
   - Evidence from document
   - Required action

3. MAJOR ISSUES (same format)

4. MINOR ISSUES (same format)

5. COMPLIANT ELEMENTS (what passed)

6. RECOMMENDATIONS
   - Immediate actions required
   - Priority order

7. COMPLIANCE CHECKLIST
   - Each rule with PASS/FAIL/PARTIAL status

Format as structured markdown."""
        
        report = self.call_llm(system_prompt, user_prompt, temperature=0.3)
        
        return {
            "report_text": report,
            "phase_results": phase_results,
            "metadata": self.metadata
        }
    
    # ========================================================================
    # MAIN ANALYSIS WORKFLOW
    # ========================================================================
    
    def analyze(self, document_path: str, prospectus_path: str, 
                rules_path: str, metadata_path: str) -> Dict:
        """
        Main analysis workflow - executes all 7 phases
        """
        print("\n" + "="*80)
        print("PROSPECTUS COMPLIANCE ANALYZER")
        print("="*80)
        print(f"Document: {document_path}")
        print(f"Prospectus: {prospectus_path}")
        print(f"Rules: {rules_path}")
        print(f"Metadata: {metadata_path}")
        
        # Load all files
        self.document_data = self.load_json_file(document_path)
        self.rules_data = self.load_json_file(rules_path)
        self.metadata = self.load_json_file(metadata_path)
        
        if prospectus_path and os.path.exists(prospectus_path):
            self.prospectus_text = self.load_docx_file(prospectus_path)
        
        # Execute all phases
        phase_results = []
        
        phase_results.append(self.phase1_document_understanding())
        phase_results.append(self.phase2_rules_framework())
        phase_results.append(self.phase3_field_presence_check())
        phase_results.append(self.phase4_content_quality_analysis())
        phase_results.append(self.phase5_cross_reference_analysis())
        phase_results.append(self.phase6_context_aware_evaluation())
        phase_results.append(self.phase7_prospectus_verification())
        
        # Generate final report
        final_report = self.generate_final_report(phase_results)
        
        return final_report
    
    def save_report(self, report: Dict, output_path: str):
        """Save the compliance report to file"""
        print(f"\nüíæ Saving report to {output_path}...")
        
        # Save as text file
        txt_path = output_path.replace('.md', '.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(report["report_text"])
        
        print(f"‚úÖ Report saved to {txt_path}")


def main():
    """Main entry point"""
    if len(sys.argv) < 5:
        print("Usage: python prospectus_compliance_analyzer.py <document.json> <prospectus.docx> <rules.json> <metadata.json>")
        print("\nExample:")
        print("  python prospectus_compliance_analyzer.py exemple.json prospectus.docx prospectus_rules.json metadata.json")
        sys.exit(1)
    
    document_path = sys.argv[1]
    prospectus_path = sys.argv[2]
    rules_path = sys.argv[3]
    metadata_path = sys.argv[4]
    
    # Get API key from .env file
    api_key = os.getenv('TOKENFACTORY_API_KEY')
    if not api_key:
        print("‚ùå Error: TOKENFACTORY_API_KEY not found in .env file")
        print("   Please create a .env file with: TOKENFACTORY_API_KEY=your-api-key")
        sys.exit(1)
    
    # Initialize analyzer
    analyzer = ProspectusComplianceAnalyzer(api_key)
    
    # Run analysis
    try:
        report = analyzer.analyze(document_path, prospectus_path, rules_path, metadata_path)
        
        # Save report
        output_path = "prospectus_validation_report.txt"
        analyzer.save_report(report, output_path)
        
        # Generate violation annotations JSON for highlighting
        annotations = generate_prospectus_violation_annotations(report, analyzer.document_data)
        annotations_file = "prospectus_violation_annotations.json"
        with open(annotations_file, 'w', encoding='utf-8') as f:
            json.dump(annotations, f, indent=2, ensure_ascii=False)
        
        print(f"üíæ Violation annotations saved to: {annotations_file}")
        
        print("\n" + "="*80)
        print("‚úÖ ANALYSIS COMPLETE")
        print("="*80)
        print(f"\nReport available at: {output_path}")
        
    except Exception as e:
        print(f"\n‚ùå Error during analysis: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


def generate_prospectus_violation_annotations(report, document):
    """
    Generate a JSON file with violation annotations for highlighting in the document.
    Parses the report text to extract violations.
    """
    annotations = {
        "document_annotations": [],
        "summary": {
            "total_violations": 0,
            "critical_violations": 0,
            "major_violations": 0,
            "minor_violations": 0
        },
        "instructions": "Use this file to highlight violations in the document. Each annotation contains the exact location, phrase, and violation details for commenting."
    }
    
    report_text = report.get("report_text", "")
    
    # Parse violations from the report text
    # Look for sections like "CRITICAL VIOLATIONS", "MAJOR ISSUES", "MINOR ISSUES"
    
    # Extract critical violations
    critical_section = re.search(r'CRITICAL VIOLATIONS?(.*?)(?:MAJOR|MINOR|COMPLIANT|RECOMMENDATIONS|$)', report_text, re.DOTALL | re.IGNORECASE)
    if critical_section:
        violations = parse_violations_from_text(critical_section.group(1), 'critical', document)
        annotations["document_annotations"].extend(violations)
        annotations["summary"]["critical_violations"] = len(violations)
        annotations["summary"]["total_violations"] += len(violations)
    
    # Extract major issues
    major_section = re.search(r'MAJOR ISSUES?(.*?)(?:MINOR|COMPLIANT|RECOMMENDATIONS|$)', report_text, re.DOTALL | re.IGNORECASE)
    if major_section:
        violations = parse_violations_from_text(major_section.group(1), 'major', document)
        annotations["document_annotations"].extend(violations)
        annotations["summary"]["major_violations"] = len(violations)
        annotations["summary"]["total_violations"] += len(violations)
    
    # Extract minor issues
    minor_section = re.search(r'MINOR ISSUES?(.*?)(?:COMPLIANT|RECOMMENDATIONS|$)', report_text, re.DOTALL | re.IGNORECASE)
    if minor_section:
        violations = parse_violations_from_text(minor_section.group(1), 'minor', document)
        annotations["document_annotations"].extend(violations)
        annotations["summary"]["minor_violations"] = len(violations)
        annotations["summary"]["total_violations"] += len(violations)
    
    # Also check phase 3 results for field presence issues
    phase_results = report.get("phase_results", [])
    for phase in phase_results:
        if phase.get("phase") == "Phase 3: Field Presence Check":
            for finding in phase.get("findings", []):
                if "EMPTY" in finding.get("result", "") or "MISSING" in finding.get("result", ""):
                    rule_id = finding.get("rule_id", "UNKNOWN")
                    severity = finding.get("severity", "major")
                    
                    annotation = {
                        "rule_id": rule_id,
                        "severity": severity,
                        "location": "document-wide",
                        "page_number": get_slide_number_from_location_prosp("document-wide", document),
                        "exact_phrase": f"Field check: {finding.get('fields_checked', [])}",
                        "character_count": 0,
                        "violation_comment": f"[{rule_id}] {finding.get('result', 'Field missing or empty')[:200]}",
                        "required_action": "Populate the required field with appropriate data"
                    }
                    
                    annotations["document_annotations"].append(annotation)
                    annotations["summary"]["total_violations"] += 1
                    if severity == "critical":
                        annotations["summary"]["critical_violations"] += 1
                    elif severity == "major":
                        annotations["summary"]["major_violations"] += 1
                    else:
                        annotations["summary"]["minor_violations"] += 1
    
    return annotations


def parse_violations_from_text(text, severity, document):
    """Parse individual violations from a section of text."""
    violations = []
    
    # Look for rule IDs (PROSP_XXX pattern)
    rule_matches = re.finditer(r'(PROSP_\d+)[:\s]+(.*?)(?=PROSP_\d+|$)', text, re.DOTALL)
    
    for match in rule_matches:
        rule_id = match.group(1)
        violation_text = match.group(2).strip()
        
        # Extract location if mentioned
        location = "document-wide"
        location_match = re.search(r'(?:Location|Found in|Section)[:\s]+([^\n]+)', violation_text, re.IGNORECASE)
        if location_match:
            location = location_match.group(1).strip()
        
        # Extract evidence/finding
        evidence = ""
        evidence_match = re.search(r'(?:Evidence|Finding|Issue)[:\s]+([^\n]+)', violation_text, re.IGNORECASE)
        if evidence_match:
            evidence = evidence_match.group(1).strip()
        
        # Extract required action
        action = ""
        action_match = re.search(r'(?:Required Action|Action|Fix)[:\s]+([^\n]+)', violation_text, re.IGNORECASE)
        if action_match:
            action = action_match.group(1).strip()
        
        annotation = {
            "rule_id": rule_id,
            "severity": severity,
            "location": location,
            "page_number": get_slide_number_from_location_prosp(location, document),
            "exact_phrase": evidence[:500] if evidence else violation_text[:200],
            "character_count": len(evidence) if evidence else len(violation_text[:200]),
            "violation_comment": f"[{rule_id}] {violation_text[:300]}",
            "required_action": action if action else "Review and correct violation"
        }
        
        violations.append(annotation)
    
    return violations


def get_slide_number_from_location_prosp(location_string, document):
    """Get the actual slide number from the document structure based on location string."""
    location_map = {
        'page_de_garde': 'page_de_garde',
        'cover': 'page_de_garde',
        'slide_2': 'slide_2',
        'page_de_fin': 'page_de_fin',
        'back page': 'page_de_fin',
        'end': 'page_de_fin',
        'fund_characteristics': 'page_de_fin'
    }
    
    location_lower = location_string.lower().strip()
    
    # For document-wide violations, try to infer from context
    if 'document-wide' in location_lower or 'entire document' in location_lower:
        # Risk-related violations typically on cover or slide 2
        if 'risk' in location_lower or 'sri' in location_lower:
            return 1
        # Characteristics typically on last page
        if 'characteristic' in location_lower or 'fee' in location_lower or 'benchmark' in location_lower:
            return document.get('page_de_fin', {}).get('slide_number', document.get('document_metadata', {}).get('page_count', 6))
        # Default to first page
        return 1
    
    # Try direct mapping first
    for key, section_name in location_map.items():
        if key in location_lower:
            if section_name in document:
                slide_num = document[section_name].get('slide_number')
                if slide_num:
                    return slide_num
    
    # Check in pages_suivantes array
    if 'pages_suivantes' in document:
        for page in document['pages_suivantes']:
            slide_num = page.get('slide_number')
            if f'slide_{slide_num}' in location_lower or f'slide {slide_num}' in location_lower:
                return slide_num
    
    # Try to extract slide number from pattern
    match = re.search(r'(?:slide|page)[_\s]?(\d+)', location_lower)
    if match:
        return int(match.group(1))
    
    # Default to first page
    return 1


if __name__ == "__main__":
    main()
