#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Pipeline Complet - Extraction et V√©rification de Conformit√©
Usage: python pipeline.py <presentation.pptx> [metadata.json] [prospectus.docx]
"""

import sys
import os
import json
import io
import time
from pathlib import Path

# Fix encoding for Windows
if sys.platform == 'win32':
    import locale
    try:
        locale.setlocale(locale.LC_ALL, 'fr_FR.UTF-8')
    except:
        try:
            locale.setlocale(locale.LC_ALL, '')
        except:
            pass
    
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', line_buffering=True)
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace', line_buffering=True)
else:
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

os.environ['PYTHONIOENCODING'] = 'utf-8'

from dotenv import load_dotenv

# Load environment
load_dotenv()

if not os.path.exists('.env'):
    print("ERROR: .env file not found! Run: python setup.py")
    sys.exit(1)

# Import extraction class from test.py
from test import PPTXFinancialExtractor

# Load the agent
# NOTE: agent_local.py already extracts prospectus.docx automatically on load
# (lines 428-556 in agent_local.py)
print("Loading agent...")
with open('agent_local.py', encoding='utf-8') as f:
    agent_code = f.read()
    exec(agent_code, globals())

# The prospectus extraction is already done by agent_local.py when it loads
# We just need to check if it was successful
def check_prospectus_extracted():
    """
    Check if prospectus was already extracted by agent_local.py
    Returns the prospectus_data if available
    """
    # prospectus_data is a global variable set by agent_local.py
    if 'prospectus_data' in globals() and globals()['prospectus_data']:
        return globals()['prospectus_data']
    return None


def extract_prospectus(prospectus_path):
    """
    Extract and parse prospectus.docx
    
    Args:
        prospectus_path: Path to prospectus.docx
    
    Returns:
        dict: Parsed prospectus data or None
    """
    if not os.path.exists(prospectus_path):
        print(f"‚ö†Ô∏è  {prospectus_path} not found - skipping prospectus checks")
        return None
    
    print(f"\n{'='*70}")
    print(f"üìÑ EXTRACTION DU PROSPECTUS")
    print(f"{'='*70}")
    
    try:
        from docx import Document
        doc_prospectus = Document(prospectus_path)
        
        # Extract text sections
        prospectus_raw_text = '\n'.join([para.text for para in doc_prospectus.paragraphs if para.text.strip()])
        
        # Also extract text from tables
        for table in doc_prospectus.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        prospectus_raw_text += '\n' + cell.text
        
        char_count = len(prospectus_raw_text)
        estimated_tokens = estimate_tokens(prospectus_raw_text)
        
        print(f"‚úì Fichier charg√©: {prospectus_path}")
        print(f"  Caract√®res: {char_count:,} (~{estimated_tokens:,} tokens)")
        
        # Try to parse with AI (Token Factory / hosted Llama)
        if tokenfactory_client and len(prospectus_raw_text) > 100:
            use_tokenfactory = True
            api_name = "Token Factory (Llama-3.1-70B)"
            
            print(f"\nü§ñ Analyse avec {api_name}...")
            
            try:
                MAX_CHARS = 80000
                
                if char_count > MAX_CHARS:
                    print(f"   Document volumineux. D√©coupage en chunks...")
                    chunks = split_document_intelligently(prospectus_raw_text, MAX_CHARS)
                    print(f"   D√©coup√© en {len(chunks)} chunks.")
                    
                    parsed_results = []
                    for i, chunk in enumerate(chunks):
                        chunk_tokens = estimate_tokens(chunk)
                        print(f"\n   Chunk {i+1}/{len(chunks)}: ~{chunk_tokens:,} tokens", end=" ")
                        
                        result_text = parse_chunk_with_ai(chunk, i, len(chunks), use_tokenfactory)
                        if result_text:
                            try:
                                cleaned = clean_json_response(result_text)
                                parsed = json.loads(cleaned)
                                parsed_results.append(parsed)
                                print("‚úì")
                            except json.JSONDecodeError as e:
                                print(f"‚ö†Ô∏è  Erreur JSON: {e}")
                                continue
                        
                        time.sleep(1)
                    
                    if parsed_results:
                        print("\n   Fusion des r√©sultats...")
                        prospectus_data = merge_parsed_results(parsed_results)
                        prospectus_data['raw_text'] = prospectus_raw_text
                        
                        extracted_fields = sum(1 for k, v in prospectus_data.items()
                                             if k != 'raw_text' and v is not None and v != {} and v != "")
                        
                        print(f"\n‚úì Prospectus analys√© ({extracted_fields}/{len(EXPECTED_FIELDS)} champs extraits)")
                    else:
                        print("\n‚ö†Ô∏è  √âchec de l'analyse - utilisation du texte brut uniquement")
                        prospectus_data = {'raw_text': prospectus_raw_text}
                
                else:
                    print("   Taille acceptable. Traitement en une seule requ√™te...")
                    result_text = parse_chunk_with_ai(prospectus_raw_text, 0, 1, use_tokenfactory)
                    
                    if result_text:
                        try:
                            cleaned = clean_json_response(result_text)
                            prospectus_data = json.loads(cleaned)
                            prospectus_data['raw_text'] = prospectus_raw_text
                            
                            extracted_fields = sum(1 for k, v in prospectus_data.items()
                                                 if k != 'raw_text' and v is not None and v != {} and v != "")
                            
                            print(f"\n‚úì Prospectus analys√© ({extracted_fields}/{len(EXPECTED_FIELDS)} champs extraits)")
                        except json.JSONDecodeError as e:
                            print(f"\n‚ö†Ô∏è  Erreur JSON: {e}")
                            prospectus_data = {'raw_text': prospectus_raw_text}
                    else:
                        print("\n‚ö†Ô∏è  Aucune r√©ponse de l'IA - utilisation du texte brut uniquement")
                        prospectus_data = {'raw_text': prospectus_raw_text}
                
                # Display extracted information
                if prospectus_data and prospectus_data != {'raw_text': prospectus_raw_text}:
                    print(f"\nüìä Informations extraites:")
                    if prospectus_data.get('fund_name'):
                        print(f"  - Fonds: {prospectus_data['fund_name']}")
                    if prospectus_data.get('sri'):
                        print(f"  - SRI: {prospectus_data['sri']}")
                    if prospectus_data.get('benchmark'):
                        benchmark = prospectus_data['benchmark']
                        print(f"  - Benchmark: {benchmark[:80]}{'...' if len(benchmark) > 80 else ''}")
                    if prospectus_data.get('management_fees'):
                        print(f"  - Frais de gestion: {prospectus_data['management_fees']}")
                    if prospectus_data.get('investment_objective'):
                        obj = prospectus_data['investment_objective']
                        print(f"  - Objectif: {obj[:80]}{'...' if len(obj) > 80 else ''}")
                
                return prospectus_data
                
            except KeyboardInterrupt:
                print(f"\n\n‚ö†Ô∏è  Analyse interrompue par l'utilisateur")
                return {'raw_text': prospectus_raw_text}
            except Exception as e:
                print(f"\n‚ö†Ô∏è  Erreur d'analyse IA: {type(e).__name__}: {str(e)}")
                return {'raw_text': prospectus_raw_text}
        else:
            return {'raw_text': prospectus_raw_text}
    
    except Exception as e:
        print(f"\n‚ùå Impossible de charger le prospectus: {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def load_metadata(metadata_path):
    """
    Load metadata.json
    
    Args:
        metadata_path: Path to metadata.json
    
    Returns:
        dict: Metadata or None
    """
    if not os.path.exists(metadata_path):
        print(f"‚ö†Ô∏è  {metadata_path} non trouv√© - utilisation des valeurs par d√©faut")
        return None
    
    try:
        with open(metadata_path, 'r', encoding='utf-8') as f:
            metadata = json.load(f)
        print(f"‚úì M√©tadonn√©es charg√©es: {metadata_path}")
        return metadata
    except Exception as e:
        print(f"‚ö†Ô∏è  Impossible de charger {metadata_path}: {e}")
        return None


def merge_metadata_with_doc(doc, metadata):
    """
    Merge metadata.json into document_metadata
    
    Args:
        doc: Document JSON structure
        metadata: Metadata dict from metadata.json
    
    Returns:
        dict: Updated document
    """
    if not metadata:
        return doc
    
    if 'document_metadata' not in doc:
        doc['document_metadata'] = {}
    
    # Map metadata fields to document_metadata
    if 'Le client est-il un professionnel' in metadata:
        is_professional = metadata['Le client est-il un professionnel']
        doc['document_metadata']['client_type'] = 'professional' if is_professional else 'retail'
        print(f"  ‚Üí Type client: {'professionnel' if is_professional else 'retail'}")
    
    if 'Soci√©t√© de Gestion' in metadata:
        doc['document_metadata']['management_company'] = metadata['Soci√©t√© de Gestion']
        print(f"  ‚Üí Soci√©t√© de gestion: {metadata['Soci√©t√© de Gestion']}")
    
    return doc


def extract_pptx(pptx_path, api_key):
    """
    Extract data from PPTX file using test.py
    IMPORTANT: Always extracts, even if a JSON file already exists
    (User may have updated the PPTX file)
    
    Args:
        pptx_path: Path to PPTX file
        api_key: Token Factory API key
    
    Returns:
        str: Path to extracted JSON file
    """
    print(f"\n{'='*70}")
    print(f"üìä EXTRACTION DE LA PR√âSENTATION PPTX")
    print(f"{'='*70}")
    
    # Generate output JSON filename based on PPTX name
    pptx_name = Path(pptx_path).stem
    output_json = f"extracted_data_{pptx_name}.json"
    
    # Check if file already exists (inform user but still extract)
    if os.path.exists(output_json):
        print(f"‚ö†Ô∏è  Fichier JSON existant d√©tect√©: {output_json}")
        print(f"   ‚Üí Nouvelle extraction va √©craser l'ancien fichier")
    
    # Always extract (even if JSON exists) - user may have updated the PPTX
    print(f"üìÑ Extraction du fichier: {pptx_path}")
    extractor = PPTXFinancialExtractor(api_key=api_key)
    extractor.process(pptx_path, output_json)
    
    return output_json


def check_document_compliance(json_file_path, prospectus_data=None):
    """
    Full document compliance checker (from check.py)
    
    Args:
        json_file_path: Path to JSON document
        prospectus_data: Parsed prospectus data (optional)
    
    Returns:
        dict with violations list
    """
    try:
        # Load document
        with open(json_file_path, 'r', encoding='utf-8') as f:
            doc = json.load(f)

        violations = []

        # Extract metadata
        doc_metadata = doc.get('document_metadata', {})
        fund_isin = doc_metadata.get('fund_isin')
        client_type = doc_metadata.get('client_type', 'retail')
        doc_type = doc_metadata.get('document_type', 'fund_presentation')
        fund_status = doc_metadata.get('fund_status', 'active')
        esg_classification = doc_metadata.get('fund_esg_classification', 'other')
        country_code = doc_metadata.get('country_code', None)
        fund_age_years = doc_metadata.get('fund_age_years', None)

        print(f"\n{'='*70}")
        print(f"üîç RAPPORT DE CONFORMIT√â")
        print(f"{'='*70}")
        print(f"Fichier: {json_file_path}")
        print(f"ISIN: {fund_isin or 'Non sp√©cifi√©'}")
        print(f"Type client: {client_type.upper()}")
        print(f"Type document: {doc_type}")
        print(f"Statut fonds: {fund_status}")
        print(f"Classification ESG: {esg_classification}")
        print(f"{'='*70}\n")
        
        # CHECK 1: REGISTRATION
        if fund_isin and funds_db:
            fund_info = None
            for fund in funds_db:
                if fund['isin'] == fund_isin:
                    fund_info = fund
                    break

            if fund_info:
                authorized_countries = fund_info['authorized_countries']
                print(f" Fonds: {fund_info['fund_name']}")
                print(f" Autoris√© dans {len(authorized_countries)} pays\n")

                reg_violations = check_registration_rules_enhanced(
                    doc,
                    fund_isin,
                    authorized_countries
                )
                violations.extend(reg_violations)

                if not reg_violations:
                    print(" ‚úÖ Conformit√© registration: OK\n")
            else:
                print(f"‚ö†Ô∏è  Fonds {fund_isin} non trouv√© dans la base de donn√©es\n")

        # CHECK 2: DISCLAIMERS
        if disclaimers_db:
            doc_type_mapping = {
                'fund_presentation': 'OBAM Presentation',
                'commercial_doc': 'Commercial documentation',
                'fact_sheet': 'OBAM Presentation'
            }

            disclaimer_type = doc_type_mapping.get(doc_type, 'OBAM Presentation')

            if disclaimer_type in disclaimers_db:
                client_key = 'professional' if client_type.lower() == 'professional' else 'retail'
                required_disclaimer = disclaimers_db[disclaimer_type].get(client_key)

                if required_disclaimer and len(required_disclaimer) > 50:
                    all_slides_text = extract_all_text_from_doc(doc)

                    if tokenfactory_client:
                        result = check_disclaimer_in_document(all_slides_text, required_disclaimer)

                        if result.get('status') == 'MISSING':
                            violations.append({
                                'type': 'DISCLAIMER',
                                'severity': 'CRITICAL',
                                'slide': 'Document-wide',
                                'location': 'Missing',
                                'rule': f'Required {client_key} disclaimer',
                                'message': f"Required disclaimer not found",
                                'evidence': result.get('explanation', ''),
                                'confidence': result.get('confidence', 0)
                            })

        # CHECK 3: STRUCTURE
        if structure_rules:
            print("üîç V√©rification structure...")
            try:
                structure_violations = check_structure_rules_enhanced(doc, client_type, fund_status)
                violations.extend(structure_violations)
                if not structure_violations:
                    print(" ‚úÖ Structure: OK\n")
            except Exception as e:
                print(f"‚ö†Ô∏è  Erreur v√©rification structure: {e}\n")

        # CHECK 4: GENERAL RULES
        if general_rules:
            print("üîç V√©rification r√®gles g√©n√©rales...")
            try:
                gen_violations = check_general_rules_enhanced(doc, client_type, country_code)
                violations.extend(gen_violations)
                if not gen_violations:
                    print(" ‚úÖ R√®gles g√©n√©rales: OK\n")
            except Exception as e:
                print(f"‚ö†Ô∏è  Erreur v√©rification r√®gles g√©n√©rales: {e}\n")

        # CHECK 5: VALUES/SECURITIES
        if values_rules:
            print("üîç V√©rification valeurs/titres...")
            try:
                values_violations = check_values_rules_enhanced(doc)
                violations.extend(values_violations)
                if not values_violations:
                    print(" ‚úÖ Valeurs/Titres: OK\n")
            except Exception as e:
                print(f"‚ö†Ô∏è  Erreur v√©rification valeurs: {e}\n")

        # CHECK 6: ESG
        if esg_rules:
            print("üîç V√©rification ESG...")
            try:
                esg_violations = check_esg_rules_enhanced(doc, esg_classification, client_type)
                violations.extend(esg_violations)
                if not esg_violations:
                    print(" ‚úÖ ESG: OK\n")
            except Exception as e:
                print(f"‚ö†Ô∏è  Erreur v√©rification ESG: {e}\n")

        # CHECK 7: PERFORMANCE
        if performance_rules:
            print("üîç V√©rification performance...")
            try:
                perf_violations = check_performance_rules_enhanced(doc, client_type, fund_age_years)
                violations.extend(perf_violations)
                if not perf_violations:
                    print(" ‚úÖ Performance: OK\n")
            except Exception as e:
                print(f"‚ö†Ô∏è  Erreur v√©rification performance: {e}\n")

        # CHECK 8: PROSPECTUS
        if prospectus_data and prospectus_rules:
            print("üîç V√©rification conformit√© prospectus...")
            try:
                prosp_violations = check_prospectus_compliance(doc, prospectus_data)
                violations.extend(prosp_violations)
                if not prosp_violations:
                    print(" ‚úÖ Prospectus: OK\n")
            except Exception as e:
                print(f"‚ö†Ô∏è  Erreur v√©rification prospectus: {e}\n")

        # FINAL REPORT
        print(f"\n{'='*70}")
        if len(violations) == 0:
            print("‚úÖ AUCUNE VIOLATION - Document conforme!")
        else:
            print(f"‚ö†Ô∏è  {len(violations)} VIOLATION(S) D√âTECT√âE(S)")
        print(f"{'='*70}\n")

        # Display violations
        for i, v in enumerate(violations, 1):
            print(f"{'='*70}")
            print(f"[{v['severity']}] {v['type']} Violation #{i}")
            print(f"{'='*70}")
            print(f"üìã R√®gle: {v['rule']}")
            print(f"‚ö†Ô∏è  Probl√®me: {v['message']}")
            print(f"üìç Localisation: {v['slide']} - {v['location']}")
            print(f"\nüìÑ Preuve:")
            print(f"   {v['evidence']}")
            print()

        # Summary
        if violations:
            print(f"\n{'='*70}")
            print(f"R√âSUM√â")
            print(f"{'='*70}")

            # By type
            type_counts = {}
            for v in violations:
                vtype = v['type']
                type_counts[vtype] = type_counts.get(vtype, 0) + 1

            print(f"\nViolations par type:")
            for vtype, count in sorted(type_counts.items()):
                print(f"   {vtype}: {count}")

            # By severity
            severity_counts = {}
            for v in violations:
                sev = v['severity']
                severity_counts[sev] = severity_counts.get(sev, 0) + 1

            print(f"\nViolations par s√©v√©rit√©:")
            for sev in ['CRITICAL', 'MAJOR', 'WARNING']:
                if sev in severity_counts:
                    print(f"   {sev}: {severity_counts[sev]}")

        return {
            'total_violations': len(violations),
            'violations': violations,
            'fund_isin': fund_isin,
            'client_type': client_type
        }

    except FileNotFoundError:
        print(f"\n‚ùå Fichier '{json_file_path}' non trouv√©")
        return {'error': 'File not found'}
    except json.JSONDecodeError as e:
        print(f"\n‚ùå Fichier JSON invalide: {e}")
        return {'error': 'Invalid JSON'}
    except Exception as e:
        print(f"\n‚ùå Erreur lors de la v√©rification: {e}")
        import traceback
        traceback.print_exc()
        return {'error': str(e)}


def main():
    """Main function"""
    if len(sys.argv) < 2:
        print("\n" + "="*70)
        print("PIPELINE COMPLET - EXTRACTION ET V√âRIFICATION DE CONFORMIT√â")
        print("="*70)
        print("\nUsage:")
        print("  python pipeline.py <presentation.pptx> <metadata.json> <prospectus.docx>")
        print("\nExemples:")
        print("  python pipeline.py presentation.pptx metadata.json prospectus.docx")
        print("  python pipeline.py presentation.pptx metadata.json")
        print("  python pipeline.py presentation.pptx")
        print("\nLe pipeline va:")
        print("  1. Extraire les donn√©es du PPTX")
        print("  2. Charger les m√©tadonn√©es (si fourni)")
        print("  3. Extraire le prospectus (si fourni)")
        print("  4. Effectuer la v√©rification de conformit√© compl√®te")
        print("\nNote: Tous les fichiers doivent √™tre dans le r√©pertoire courant")
        print("="*70)
        sys.exit(1)

    # R√©cup√©rer les fichiers en entr√©e
    pptx_file = sys.argv[1]
    metadata_file = sys.argv[2] if len(sys.argv) > 2 else None
    prospectus_file = sys.argv[3] if len(sys.argv) > 3 else None
    
    # V√©rifier que le fichier PPTX existe
    if not os.path.exists(pptx_file):
        print(f"\n‚ùå Fichier '{pptx_file}' non trouv√©")
        print(f"üìÅ R√©pertoire actuel: {os.getcwd()}")
        sys.exit(1)
    
    # V√©rifier les extensions
    if not pptx_file.lower().endswith('.pptx'):
        print(f"\n‚ùå Le fichier '{pptx_file}' doit √™tre un fichier .pptx")
        sys.exit(1)
    
    if metadata_file and not metadata_file.lower().endswith('.json'):
        print(f"\n‚ö†Ô∏è  Attention: '{metadata_file}' devrait √™tre un fichier .json")
    
    if prospectus_file and not prospectus_file.lower().endswith('.docx'):
        print(f"\n‚ö†Ô∏è  Attention: '{prospectus_file}' devrait √™tre un fichier .docx")
    
    # Get API key
    api_key = os.getenv('TOKENFACTORY_API_KEY')
    if not api_key:
        print("ERROR: TOKENFACTORY_API_KEY not found in .env file")
        sys.exit(1)
    
    print("\n" + "="*70)
    print("üöÄ D√âMARRAGE DU PIPELINE")
    print("="*70)
    print(f"üìÑ Pr√©sentation: {pptx_file}")
    if metadata_file:
        print(f"üìã M√©tadonn√©es: {metadata_file}")
    else:
        print(f"üìã M√©tadonn√©es: Non fourni")
    if prospectus_file:
        print(f"üìë Prospectus: {prospectus_file}")
    else:
        print(f"üìë Prospectus: Non fourni")
    print("="*70)
    
    # V√©rifier que les fichiers fournis existent
    if metadata_file and not os.path.exists(metadata_file):
        print(f"\n‚ö†Ô∏è  Fichier m√©tadonn√©es '{metadata_file}' non trouv√© - ignor√©")
        metadata_file = None
    
    if prospectus_file and not os.path.exists(prospectus_file):
        print(f"\n‚ö†Ô∏è  Fichier prospectus '{prospectus_file}' non trouv√© - ignor√©")
        prospectus_file = None
    
    # Step 1: Extract PPTX
    # IMPORTANT: Always extracts, even if JSON already exists (user may have updated PPTX)
    json_file = extract_pptx(pptx_file, api_key)
    
    # Step 2: Load and merge metadata
    # IMPORTANT: Always loads the provided metadata file (user may have updated it)
    if metadata_file:
        metadata = load_metadata(metadata_file)
        if metadata:
            print(f"\nüìã Fusion des m√©tadonn√©es...")
            with open(json_file, 'r', encoding='utf-8') as f:
                doc = json.load(f)
            doc = merge_metadata_with_doc(doc, metadata)
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(doc, f, indent=2, ensure_ascii=False)
    else:
        print(f"\n‚ö†Ô∏è  Aucun fichier m√©tadonn√©es fourni - utilisation des valeurs par d√©faut")
    
    # Step 3: Extract prospectus
    # IMPORTANT: Si un fichier prospectus est fourni, TOUJOURS l'extraire (nouveau)
    # On IGNORE compl√®tement celui extrait par agent_local.py si un fichier est fourni
    # Car l'utilisateur peut fournir un nouveau prospectus √† chaque ex√©cution
    prospectus_data = None  # Initialiser √† None pour forcer l'extraction si fichier fourni
    
    if prospectus_file:
        # Fichier prospectus fourni - TOUJOURS extraire (m√™me si agent_local.py en a d√©j√† extrait un)
        print(f"\nüìÑ Extraction du prospectus FOURNI: {prospectus_file}")
        print(f"   ‚Üí Ignore l'ancien prospectus extrait par agent_local.py")
        prospectus_data = extract_prospectus(prospectus_file)
        if prospectus_data:
            # Mettre √† jour la variable globale pour √©craser l'ancien prospectus
            globals()['prospectus_data'] = prospectus_data
            print(f"‚úì Prospectus NOUVEAU extrait et analys√©")
            print(f"   ‚Üí Ce prospectus sera utilis√© pour les v√©rifications")
        else:
            print(f"‚ö†Ô∏è  √âchec de l'extraction du prospectus")
            prospectus_data = None
    else:
        # Aucun fichier prospectus fourni - utiliser celui extrait par agent_local.py (si disponible)
        prospectus_data = check_prospectus_extracted()
        if prospectus_data:
            print(f"\n‚úì Utilisation du prospectus d√©j√† extrait par agent_local.py (prospectus.docx)")
            print(f"   Note: Pour utiliser un NOUVEAU prospectus, fournissez-le en param√®tre")
        else:
            print(f"\n‚ö†Ô∏è  Aucun prospectus fourni - les v√©rifications prospectus seront ignor√©es")
            prospectus_data = None
    
    # Step 5: Run compliance check
    result = check_document_compliance(json_file, prospectus_data)
    
    if 'error' not in result:
        print(f"\n{'='*70}")
        print("‚úÖ V√âRIFICATION TERMIN√âE")
        print(f"{'='*70}")
        
        if result['total_violations'] == 0:
            print("\nüéâ Document enti√®rement conforme!")
            sys.exit(0)
        else:
            print(f"\n‚ö†Ô∏è  Veuillez examiner et corriger {result['total_violations']} violation(s)")
            sys.exit(1)
    else:
        sys.exit(1)


if __name__ == "__main__":
    main()

